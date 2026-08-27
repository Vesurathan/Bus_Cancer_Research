"""Insert the five new figures (training/validation accuracy, training/validation
loss, confusion matrix, precision-recall, F1-vs-threshold) into the user's
existing dissertation .docx, matching the document's figure formatting.
Writes to a NEW file so the original is untouched."""
import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL

SRC = "../doc/Dissertation_Improved_15500_Plus.docx"
DST = "../doc/Dissertation_Improved_with_figures.docx"
FIGDIR = "../dissertation/figures"

d = docx.Document(SRC)

# locate the "4.4  Comparison with Existing Methods" heading to insert before it
anchor = None
for p in d.paragraphs:
    if p.style.name.startswith("Heading") and "Comparison with Existing" in p.text:
        anchor = p
        break
if anchor is None:
    raise SystemExit("anchor heading '4.4 Comparison...' not found")


def add_blank():
    anchor.insert_paragraph_before("")


def add_subhead(text):
    p = anchor.insert_paragraph_before()
    r = p.add_run(text); r.bold = True


def add_image(path, width_in):
    p = anchor.insert_paragraph_before()
    p.alignment = AL.CENTER
    p.add_run().add_picture(path, width=Inches(width_in))


def add_caption(text):
    p = anchor.insert_paragraph_before()
    p.alignment = AL.CENTER
    r = p.add_run(text); r.bold = True


def add_body(text):
    p = anchor.insert_paragraph_before()
    p.alignment = AL.JUSTIFY
    p.add_run(text)


FIGS = [
    (f"{FIGDIR}/fig_train_val_accuracy.png", 5.0,
     "Figure 4.5  Training and validation accuracy of the image ViT stream over 15 epochs "
     "on an 80/20 stratified split of BrEaST.",
     "Figure 4.5 traces the image-only Vision Transformer as it learns. Training accuracy climbs "
     "steadily to roughly 0.93, but validation accuracy plateaus in the 0.65-0.75 band and peaks "
     "around epoch nine, opening a clear train-validation gap. This gap is the expected signature of "
     "over-fitting when a high-capacity backbone is fine-tuned on only around 200 training images, "
     "and it is precisely why the image stream is the weakest single evidence source (AUC 0.813) and "
     "why fusion with the descriptor and retrieval streams, together with the later addition of "
     "external data, is necessary."),
    (f"{FIGDIR}/fig_train_val_loss.png", 5.0,
     "Figure 4.6  Training and validation loss (binary cross-entropy) of the image ViT stream over "
     "the same run.",
     "Figure 4.6 shows the corresponding binary cross-entropy loss and tells the same story. "
     "Training loss falls steadily toward 0.19, whereas validation loss reaches its minimum near "
     "epoch six (about 0.80) and then drifts upward, confirming that the single-stream image model "
     "begins to over-fit early. This divergence motivates the short training horizon, the "
     "class-weighting used during optimisation, and the multi-evidence design that compensates for "
     "the image stream's limited standalone generalisation."),
    (f"{FIGDIR}/fig_confusion_matrix.png", 6.3,
     "Figure 4.7  Confusion matrices for the stacked fusion at two operating points (leak-free "
     "5-fold CV on BrEaST, n = 256): (a) the balanced 0.5 threshold and (b) the sensitivity-first "
     "0.32 threshold.",
     "Figure 4.7 contrasts the fusion's error structure at two operating points. At the balanced 0.5 "
     "threshold (panel a) the model makes 12 false negatives and 24 false positives; lowering the "
     "threshold to the sensitivity-first 0.32 (panel b) cuts false negatives to just three "
     "(sensitivity 0.97) at the cost of a modest rise in false positives to 32 (specificity 0.80). "
     "Because a missed malignancy is far costlier than a benign lesion sent for further work-up, this "
     "trade is clinically appropriate. Both errors cannot be driven to zero simultaneously: the "
     "fewest total errors achievable at any threshold is 35 of 256, a direct consequence of the "
     "0.932 area under the ROC curve, and a perfectly clean matrix on held-out data would in fact "
     "signal leakage rather than success."),
    (f"{FIGDIR}/fig_pr_curve.png", 5.0,
     "Figure 4.8  Precision-recall curves for the evidence streams and their fusion (internal "
     "5-fold CV), with average precision (AP) shown in the legend.",
     "Figure 4.8 reports precision-recall behaviour, which is more informative than ROC under class "
     "imbalance. The stacked fusion attains the highest average precision (AP = 0.890), ahead of the "
     "descriptor (0.858), image ViT (0.730) and k-NN (0.624) streams, mirroring the ranking seen in "
     "the ROC and AUC analysis."),
    (f"{FIGDIR}/fig_f1_threshold.png", 5.0,
     "Figure 4.9  F1, precision and recall of the fused model as a function of the decision "
     "threshold.",
     "Figure 4.9 sweeps the decision threshold. F1 is maximised near 0.32, but because a missed "
     "malignancy is far costlier than a false alarm the deployed operating point is deliberately set "
     "lower to favour recall, trading a little precision for sensitivity as described above."),
]

add_subhead("Learning dynamics and additional diagnostic plots")
add_body(
    "To document the model's training behaviour and error structure in more detail, this subsection "
    "reports the image stream's learning curves together with the fused model's confusion matrix, "
    "precision-recall curves and threshold sensitivity. The learning curves are taken from a single "
    "representative run on an 80/20 stratified split of BrEaST and are informative in their own "
    "right: they expose the over-fitting of the image-only stream that the multi-evidence fusion is "
    "designed to counteract. The receiver-operating-characteristic (AUC) curves were presented "
    "earlier as Figure 4.1.")
for path, w, cap, body in FIGS:
    add_blank()
    add_image(path, w)
    add_caption(cap)
    add_body(body)

d.save(DST)
print("saved", DST)
